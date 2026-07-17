# -*- coding: utf-8 -*-
"""
reglement_gen.py — Remplit LE modèle officiel PersoProject de règlement de travail.

Le modèle Word (FR ou NL) est un mail-merge : ses champs ont été transformés en
jetons {{Nom_Em}}, {{Commission_paritaire_Em}}, … (voir _reglement_bundle/). On
charge le modèle depuis Supabase Storage, on remplace les jetons par les données
(BCE + formulaire), et on renvoie le .docx. Les champs inconnus deviennent un
blanc « ………… » que la gestionnaire complète.

⚠ Document = PROJET, à faire valider juridiquement avant dépôt.
"""
import io
import json
import os
import re
from docx import Document
from docx.oxml.ns import qn

BLANK = '…………'

# --- Données officielles embarquées (dossier donnees/, versionné avec le code) ----
# Source : SPF Emploi. Régénérables avec les scripts documentés dans donnees/README.md.
#   fse_{fr,nl}.json          : dénomination OFFICIELLE de chaque CP (point 9) +
#                               son fonds de sécurité d'existence (point 4)
#   institutions_controle.json: adresse des services de contrôle par province (art. 66)
_DONNEES = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'donnees')
_CACHE = {}


def _donnees(nom):
    """Charge donnees/<nom>.json (une fois). {} si le fichier manque : le règlement
    doit sortir même sans ces données, avec des blancs à compléter."""
    if nom not in _CACHE:
        try:
            with open(os.path.join(_DONNEES, f'{nom}.json'), encoding='utf-8') as f:
                _CACHE[nom] = json.load(f)
        except (OSError, ValueError) as e:
            print(f'[REGLEMENT] donnees/{nom}.json illisible : {e}')
            _CACHE[nom] = {}
    return _CACHE[nom]


def _fse_info(cp, lang):
    """{denomination, fonds:[…]} de la CP, depuis la liste officielle des Fonds de
    sécurité d'existence (SPF Emploi). {} si la CP n'y figure pas."""
    d = _donnees('fse_nl' if str(lang).upper() == 'NL' else 'fse_fr')
    return d.get(_cp_norm(cp)) or {}


# Les fonds « 2e pilier » (pension complémentaire sectorielle) ne sont PAS ce que vise
# le point 4 « Fonds de sécurité d'existence ou Fonds social ». On les écarte par leur
# NOM et non par leur rang : la page du SPF les liste parfois EN PREMIER (FR 220,
# 327.01, 329.01, 331 ; NL 118, 220, 329.01), et l'ordre diffère même entre les listes
# FR et NL d'une même CP — se fier au rang faisait imprimer le fonds de pension, avec
# son adresse, dans un document légal. Vérifié sur les 128 CP : aucun fonds social
# n'est écarté à tort (« prépension »/« brugpensioen » n'existent dans aucun nom).
# « pensio » couvre pension, pensioen et « Pensio+ ».
RE_FONDS_PENSION = re.compile(r'pensio|pilier|pijler', re.I)


def _fonds_principal(cp, lang):
    """Le fonds à citer au point 4 : le premier fonds ADRESSABLE qui ne soit pas un
    fonds de pension « 2e pilier ».
    None si la CP n'a aucun fonds social — cas réel des SCP 102.01 et 102.09, dont le
    SPF ne liste qu'un fonds 2e pilier : mieux vaut un blanc à compléter qu'un fonds
    faux. None aussi si la CP est trop imprécise pour trancher (ex. CP 140 nue :
    chaque sous-secteur 140.01…140.09 a SON fonds)."""
    fonds = [f for f in _fse_info(cp, lang).get('fonds', []) if f.get('adresse')]
    return next((f for f in fonds if not RE_FONDS_PENSION.search(f.get('nom') or '')), None)


def _decoupe_adresse_plate(adr):
    """« Quai de Willebroeck 37, 1000 BRUXELLES » -> (rue, n°, code postal, localité)."""
    m = re.match(r'^(.*?),?\s*(\d{4})\s+([^,]+)$', (adr or '').strip())
    if not m:
        return '', '', '', ''
    voie, cp, loc = m.group(1).strip(' ,'), m.group(2), m.group(3).strip()
    mv = re.match(r'^(.*?[^\d\s])\s+(\d[\w\s/.-]*)$', voie)
    return (mv.group(1).strip(), mv.group(2).strip(), cp, loc) if mv else (voie, '', cp, loc)


# CP dont les lieux de travail changent en permanence : on n'y met PAS de siège
# d'exploitation (chantiers pour la construction, itinéraires pour le transport).
CP_SANS_SIEGE_EXPLOITATION = ('124', '140')


def _sans_siege_exploitation(cps):
    """True si l'une des CP du dossier est la construction (124) ou le transport (140),
    sous-numéros compris (124.01, 140.03, …)."""
    for cp in cps:
        base = _cp_norm(cp).split('.')[0]
        if base in CP_SANS_SIEGE_EXPLOITATION:
            return True
    return False


def _cps_du_payload(payload):
    """Toutes les CP du dossier, dans l'ordre : les régimes s'il y en a, sinon les
    champs simples ouvrier/employé."""
    regs = payload.get('regimes') if isinstance(payload.get('regimes'), list) else []
    cps = [str(r.get('cp') or '') for r in regs if isinstance(r, dict) and r.get('cp')]
    if not cps:
        cps = [str(payload.get('commission_paritaire') or ''),
               str(payload.get('cp_employe') or '')]
    return [c for c in cps if _cp_norm(c)]

JOURS_FR = ['lundi', 'mardi', 'mercredi', 'jeudi', 'vendredi', 'samedi', 'dimanche']
JOURS_NL = ['maandag', 'dinsdag', 'woensdag', 'donderdag', 'vrijdag', 'zaterdag', 'zondag']

# Personne de confiance + conseiller harcèlement selon le SEPPT (source : listes
# PersoProject — Excel « personne de confiance RT » + Attentia + Arista).
SEPPT_PDC = {
    'securex':  {'pc': 'Delphine PIETERS (0800/100.59)',
                 'harc': 'Delphine PIETERS (0800/100.59 — Health_safety@securex.be)'},
    'mensura':  {'pc': 'Benoit VAN TICHELEN et Dominique DEHON (02/549.71.07 et 011/26.40.90)',
                 'harc': 'Benoit VAN TICHELEN (02/549.71.07)'},
    'cesi':     {'pc': 'Tina SCHOLIERS (02/761.17.74)',
                 'harc': 'Mme MOREAU (02/771.00.25)'},
    'idewe':    {'pc': 'N° général (02/237.33.24)',
                 'harc': 'N° général (02/237.33.24)'},
    'proviko':  {'pc': 'Suzi Broupkmans (02/250.00.57)',
                 'harc': 'Suzi Broupkmans (02/250.00.57)'},
    'mediwet':  {'pc': 'Peter VAN SLEMBROUCK (09/221.06.07)',
                 'harc': 'Peter VAN SLEMBROUCK (09/221.06.07)'},
    'adhesia':  {'pc': 'N° général (02/511.06.86)',
                 'harc': 'N° général (02/511.06.86)'},
    'liantis':  {'pc': "Ana Rodriguez Sio — L'équipe psychosociale (078/150.888)",
                 'harc': "L'équipe psychosociale (078/150.888)"},
    'attentia': {'pc': "L'équipe des conseillers en prévention aspects psychosociaux "
                       "(02/738 75 31 — psy.prev@attentia.be)",
                 'harc': "L'équipe des conseillers en prévention aspects psychosociaux "
                         "(02/738 75 31 — psy.prev@attentia.be)"},
    'arista':   {'pc': 'Arista — conseiller psychosocial (02 533 74 11 — arista@arista.be)',
                 'harc': 'Arista — conseiller psychosocial (02 533 74 11 — arista@arista.be)'},
}


def _pdc_du_seppt(seppt):
    """Retourne {pc, harc} de la personne de confiance selon le SEPPT, ou None."""
    s = (seppt or '').lower()
    for cle, val in SEPPT_PDC.items():
        if cle in s:
            return val
    return None


def _jour(idx, lang):
    """Nom de jour localisé depuis un index 0=lundi (ou un nom déjà écrit)."""
    if idx in (None, ''):
        return BLANK
    tab = JOURS_NL if lang == 'NL' else JOURS_FR
    try:
        return tab[int(idx) % 7]
    except (ValueError, TypeError):
        return str(idx)


# Chaque institution -> les jetons du modèle (Nom/Rue/N°/CP/Localité).
#
# ⚠ Les modèles FR et NL ne numérotent PAS les jetons pareil : le même
# {{Nom_1_institution_Inst_v1}} désigne la caisse de vacances en FR et l'assurance-loi
# en NL. Une table unique (calée sur le FR) écrivait donc, dans tout règlement NL, la
# caisse dans la case assurance et le SEPPT dans la case du fonds. D'où une table PAR
# LANGUE. Relevé le 2026-07-17 directement dans les 2 .docx (Article 2 et Article 66) ;
# à revérifier si les modèles Word changent — voir tests/test_reglement_gen.py.
#
#   suffixe   |        FR (Article 2)         |        NL (Artikel 2)
#   ----------+-------------------------------+------------------------------
#   (aucun)   | 6. SEPPT                      | 2. Kas voor jaarlijkse vakantie
#   _v1       | 2. Caisse de vacances         | 3. Verzekeringsmaatschappij
#   _v2       | (annexe)                      | 4. Fondsen voor bestaanszekerheid
#   _v3       | (absent)                      | 6. Externe Dienst (SEPPT)
#   _v4       | 4. Fonds de sécurité          | 8. Kantoor voor directe belastingen
#   _v5       | 3. Assurance-loi              | (absent)
def _toks(nom, suffixe):
    s = f'_{suffixe}' if suffixe else ''
    return {'nom': f'Nom_1_institution_Inst{s}' if nom else None,
            'rue': f'Rue_institution_Inst{s}', 'no': f'No_de_la_maison_inst_Inst{s}',
            'cp': f'Code_postal_Institutions{s}', 'loc': f'Localité_institution_Inst{s}'}


def _inst_map(caisse, assurance, fonds, seppt, lois, bienetre):
    """(suffixe du nom, suffixe de l'adresse) par institution — ils diffèrent en FR."""
    m = {}
    for cle, (sn, sa) in (('caisse', caisse), ('assurance', assurance), ('fonds', fonds),
                          ('seppt', seppt)):
        t = _toks(True, sn)
        t.update({k: v for k, v in _toks(False, sa).items() if k != 'nom'})
        m[cle] = t
    for cle, sa in (('controle_lois', lois), ('controle_bienetre', bienetre)):
        t = _toks(False, sa)
        t.pop('nom')
        m[cle] = t
    return m


INST_TOKENS_FR = _inst_map(caisse=('v1', 'v4'), assurance=('v5', 'v3'), fonds=('v4', 'v5'),
                           seppt=('', 'v1'), lois='v2', bienetre='v6')
# NL : « Inspectie van de sociale wetten » = Contrôle des lois sociales (_v5),
#      « Medische inspectie » = Contrôle du bien-être au travail (_v6).
INST_TOKENS_NL = _inst_map(caisse=('', 'v1'), assurance=('v1', 'v2'), fonds=('v2', 'v3'),
                           seppt=('v3', 'v4'), lois='v5', bienetre='v6')


def _inst_tokens(lang):
    return INST_TOKENS_NL if str(lang).upper() == 'NL' else INST_TOKENS_FR


INST_TOKENS = INST_TOKENS_FR      # compat : ancien nom, modèle FR


def _province_from_cp(cp):
    try:
        n = int(re.sub(r'\D', '', str(cp))[:4])
    except (ValueError, TypeError):
        return ''
    if 1000 <= n <= 1299: return 'Bruxelles'
    if 1300 <= n <= 1499: return 'Brabant wallon'
    if 1500 <= n <= 1999 or 3000 <= n <= 3499: return 'Brabant flamand'
    if 2000 <= n <= 2999: return 'Anvers'
    if 3500 <= n <= 3999: return 'Limbourg'
    if 4000 <= n <= 4999: return 'Liège'
    if 5000 <= n <= 5999: return 'Namur'
    if 6000 <= n <= 6599 or 7000 <= n <= 7999: return 'Hainaut'
    if 6600 <= n <= 6999: return 'Luxembourg'
    if 8000 <= n <= 8999: return 'Flandre occidentale'
    if 9000 <= n <= 9999: return 'Flandre orientale'
    return ''


def _valeurs_officielles(payload, identity, lang, cps):
    """Valeurs issues des données officielles embarquées, sans le répertoire :
      - point 4  : le fonds de sécurité d'existence de la CP (nom + adresse) ;
      - article 66 : les services de contrôle compétents pour la province du client.
    Le répertoire, s'il est rempli, écrase ces valeurs (cf. _valeurs)."""
    out = {}
    idd = identity or {}
    _, _, cp_cli, _ = _decoupe_adresse(idd)
    prov = _province_from_cp(cp_cli)

    TOKS = _inst_tokens(lang)

    # --- Point 4 : Fonds de sécurité d'existence ou Fonds social
    fonds = next((_fonds_principal(c, lang) for c in cps if _fonds_principal(c, lang)), None)
    if fonds:
        rue, no, cpf, loc = _decoupe_adresse_plate(fonds.get('adresse'))
        toks = TOKS['fonds']
        out[toks['nom']] = fonds.get('nom') or ''
        out[toks['rue']], out[toks['no']] = rue, no
        out[toks['cp']], out[toks['loc']] = cpf, loc

    # --- Article 66 : Contrôle des lois sociales + Contrôle du bien-être au travail
    ctrl = _donnees('institutions_controle')
    cle = 'nl' if str(lang).upper() == 'NL' else 'fr'
    for typ in ('controle_lois', 'controle_bienetre'):
        adr = (ctrl.get(typ) or {}).get(prov, {}).get(cle)
        if not adr:
            continue
        toks = TOKS[typ]
        out[toks['rue']] = adr.get('rue') or ''
        out[toks['no']] = adr.get('no') or ''
        out[toks['cp']] = adr.get('cp') or ''
        out[toks['loc']] = adr.get('localite') or ''
    return out


def _valeurs_institutions(payload, identity, repertoire, lang='FR'):
    """Remplit les jetons d'adresses d'institutions depuis le répertoire.
    Assurance/caisse/SEPPT : par nom ; bureaux de contrôle : par province du client."""
    out = {}
    if not repertoire:
        return out
    idd = identity or {}
    _, _, cp_cli, _ = _decoupe_adresse(idd)
    prov = _province_from_cp(cp_cli)
    TOKS = _inst_tokens(lang)

    def pose(inst, typ):
        toks = TOKS.get(typ, {})
        if toks.get('nom') and inst.get('nom'):
            out[toks['nom']] = inst['nom']
        if toks.get('rue'):
            out[toks['rue']] = inst.get('rue') or ''
        if toks.get('no'):
            out[toks['no']] = inst.get('numero') or ''
        if toks.get('cp'):
            out[toks['cp']] = inst.get('code_postal') or ''
        if toks.get('loc'):
            out[toks['loc']] = inst.get('localite') or ''

    def trouve_par_nom(typ, nom):
        nom = (nom or '').lower().strip()
        cands = [i for i in repertoire if (i.get('type') or '') == typ]
        if nom:
            # 1) correspondance EXACTE prioritaire (évite qu'un nom court comme « AG »
            #    matche par erreur « AGACHE » via sous-chaîne)
            for i in cands:
                if (i.get('nom') or '').lower().strip() == nom:
                    return i
            # 2) repli : sous-chaîne dans un sens ou l'autre
            for i in cands:
                cn = (i.get('nom') or '').lower().strip()
                if cn and (cn in nom or nom in cn):
                    return i
        return cands[0] if len(cands) == 1 else None

    for typ, champ in (('assurance', 'assurance_loi'), ('caisse', 'caisse_vacances'), ('seppt', 'seppt')):
        inst = trouve_par_nom(typ, payload.get(champ))
        if inst:
            pose(inst, typ)
    # fonds : s'il n'y en a qu'un, ou celui de la province
    fonds = [i for i in repertoire if (i.get('type') or '') == 'fonds']
    if len(fonds) == 1:
        pose(fonds[0], 'fonds')
    # bureaux de contrôle : par province du client (ou l'unique)
    for typ in ('controle_lois', 'controle_bienetre'):
        cands = [i for i in repertoire if (i.get('type') or '') == typ]
        pick = next((i for i in cands if (i.get('province') or '') == prov and prov), None) \
            or (cands[0] if len(cands) == 1 else None)
        if pick:
            pose(pick, typ)
    return out


def _decoupe_adresse(idd):
    """(adresse_1, adresse_2 BCE) -> (rue, n°, code postal, localité)."""
    adr1 = (idd.get('adresse_siege_social_1') or '').strip()
    adr2 = (idd.get('adresse_siege_social_2') or '').strip()
    cp = loc = ''
    m = re.match(r'\s*(\d{4})\s+(.+)', adr2)
    if m:
        cp, loc = m.group(1), m.group(2).strip()
    rue, nomaison = adr1, ''
    m2 = re.match(r'(.+?)\s+(\d+\s*[A-Za-z]?)\s*$', adr1)
    if m2:
        rue, nomaison = m2.group(1).strip(), m2.group(2).strip()
    return rue, nomaison, cp, loc


def _valeurs(payload, identity, repertoire=None):
    """Correspondance jeton -> valeur. Ce qu'on n'a pas -> BLANK (à compléter)."""
    idd = identity or {}
    lang = 'NL' if str(payload.get('reglement_langue') or 'FR').upper() == 'NL' else 'FR'
    rue, nomaison, cp, loc = _decoupe_adresse(idd)
    ent = re.sub(r'\D', '', payload.get('num_entreprise', '') or '')
    if len(ent) == 9:
        ent = '0' + ent
    ent_fmt = f"{ent[0:4]}.{ent[4:7]}.{ent[7:10]}" if len(ent) == 10 else (ent or BLANK)
    # _cp_norm et pas re.sub(r'\D') : « 140.03 » doit rester « 140.03 », pas « 14003 ».
    cpnum = _cp_norm(payload.get('commission_paritaire', '') or '')
    cps = _cps_du_payload(payload)
    adr1 = (idd.get('adresse_siege_social_1') or '').strip()

    def ou(v):  # valeur ou blanc
        return (str(v).strip() if v not in (None, '') else BLANK)

    v = {
        # --- Société (BCE + formulaire) ---
        'Nom_Em': ou(idd.get('nom_societe')),
        'Forme_juridique_Em': ou(idd.get('forme_juridique')),
        'Adresse_Em': ou(adr1),
        'Rue_Em': ou(rue),
        'No_de_la_maison_Em': ou(nomaison),
        'Code_postal_Em': ou(cp),
        'Localité_Em': ou(loc),
        'Activité_générale_Em': ou(idd.get('secteur_activite')),
        'Téléphone_Em': ou(payload.get('telephone')),
        'No_ONSS_Em': ou(payload.get('num_onss')),
        'No_d_entreprise_Emp': ent_fmt,
        'No_employeur_Em': ou(payload.get('num_onss')),
        'No_affiliation_instit_Em': BLANK,
        'No_affiliation_instit_Em_v1': BLANK,
        'Commission_paritaire_Em': ou(cpnum),
        'Commission_paritaire_Em_v1': BLANK,
        # --- ONSS : adresse officielle (constante) ---
        'Rue_institution_Inst': 'Place Victor Horta',
        'No_de_la_maison_inst_Inst': '11',
        'Code_postal_Institutions': '1060',
        'Localité_institution_Inst': 'Bruxelles',
        # (les noms d'institutions sont posés plus bas : leurs jetons dépendent de la
        #  langue du modèle — cf. INST_TOKENS_FR / INST_TOKENS_NL)
        # --- Annexe 4 (bien-être) + Annexe 5 (lieux) ---
        'ps_nom1': ou(payload.get('premiers_soins_noms')),
        'ps_lieu1': ou(payload.get('premiers_soins_lieux')),
        'ps_nom2': BLANK, 'ps_lieu2': BLANK,
        'boite_secours': ou(payload.get('boite_secours_emplacement')),
        # Personne de confiance + harcèlement : selon le SEPPT (saisie manuelle prioritaire)
        'personne_confiance': ou(payload.get('personne_de_confiance')
                                 or (_pdc_du_seppt(payload.get('seppt')) or {}).get('pc')),
        'harcelement': ou((_pdc_du_seppt(payload.get('seppt')) or {}).get('harc')),
        # sièges d'exploitation (annexe 5) : vraies unités d'établissement BCE si
        # disponibles (identity['sieges_exploitation']), sinon repli sur le siège social
        # Construction (CP 124) et transport (CP 140) : PAS de siège d'exploitation —
        # chantiers et itinéraires changent en permanence, une adresse figée serait
        # fausse dès le lendemain.
        'sieges_exploitation': BLANK if _sans_siege_exploitation(cps) else ou(
            idd.get('sieges_exploitation')
            or ' — '.join(x for x in [adr1, (idd.get('adresse_siege_social_2') or '').strip()] if x)),
        # --- Cadre horaire (Article 10 §2) depuis les heures d'ouverture ---
        'cadre_debut': ou(payload.get('ouverture_debut')),
        'cadre_fin': ou(payload.get('ouverture_fin')),
        'cadre_jour_debut': _jour(payload.get('ouverture_jour_debut'), lang),
        'cadre_jour_fin': _jour(payload.get('ouverture_jour_fin'), lang),
        'cadre_min': '3',
        # Cadre légal (loi 1/6/2026) : min 3h/jour, max 9h/jour (le max hebdo de 50h
        # est ajouté en ligne séparée par _ajouter_max_hebdo, le modèle n'ayant pas de jeton).
        'cadre_max': '9',
    }
    # 1) noms d'institutions saisis au formulaire, dans le jeton de LA BONNE langue
    TOKS = _inst_tokens(lang)
    for typ, champ in (('assurance', 'assurance_loi'), ('caisse', 'caisse_vacances'),
                       ('seppt', 'seppt')):
        val = payload.get(champ)
        if val:
            v[TOKS[typ]['nom']] = str(val).strip()
    # 2) données officielles embarquées (fonds de la CP, services de contrôle)
    v.update({k: val for k, val in _valeurs_officielles(payload, identity, lang, cps).items() if val})
    # 3) puis le répertoire : ce que la gestionnaire a saisi elle-même fait foi.
    #    Le bloc d'une institution est repris EN ENTIER — un champ vide devient un blanc
    #    à compléter, il ne laisse PAS passer la valeur officielle. Sinon un fonds saisi
    #    sans adresse donnait un hybride : le nom saisi + l'adresse officielle d'un AUTRE
    #    organisme, ce qui est pire que l'un ou l'autre dans un document légal.
    v.update({k: (val or BLANK) for k, val in
              _valeurs_institutions(payload, identity, repertoire, lang).items()})
    # Tout jeton non couvert -> BLANK (renseigné dynamiquement au remplissage)
    return v


def _remplacer_dans_paragraphe(p, motif, remplacement):
    """Remplace `motif` dans un paragraphe MÊME s'il est éclaté sur plusieurs runs.
    Word découpe le texte au moindre changement (« (annexe n » / « ° » / « 10). ») :
    chercher run par run ne trouverait rien. On reconstitue le texte, on repère la
    correspondance, et on ne réécrit que les runs concernés — la mise en forme des
    autres est préservée. True si un remplacement a eu lieu."""
    runs = p.runs
    textes = [r.text or '' for r in runs]
    m = re.search(motif, ''.join(textes), re.I)
    if not m:
        return False
    deb, fin = m.span()
    pos, bornes = 0, []
    for t in textes:
        bornes.append((pos, pos + len(t)))
        pos += len(t)
    ecrit = False
    for k, (a, b) in enumerate(bornes):
        if b <= deb or a >= fin:
            continue
        avant = textes[k][:max(0, deb - a)]
        apres = textes[k][fin - a:] if fin <= b else ''
        runs[k].text = (avant + remplacement + apres) if not ecrit else (avant + apres)
        ecrit = True
    return True


def _corriger_renvoi_annexe(doc, lang):
    """Article 4 (Exemplaire pour le travailleur) : le modèle renvoie à l'annexe 10,
    or l'annexe 10 est « Politique de maintien du contact avec les travailleurs en
    incapacité » — l'accusé de réception dont parle l'article est l'annexe 11
    (« Accusé de réception » / « Ontvangstbewijs »). Erreur du modèle Word, corrigée
    ici à la génération (le .docx source est sur Supabase Storage, hors du dépôt).
    Ne touche NI la table des matières NI le titre de l'annexe 10, qui sont sans
    parenthèses. No-op si le modèle est corrigé un jour."""
    if str(lang).upper() == 'NL':
        motif, rempl = r'\(\s*zie\s+bijlage\s*n\s*r\.?\s*10\s*\)', '(zie bijlage nr. 11)'
    else:
        motif, rempl = r'\(\s*annexe\s*n\s*°\s*10\s*\)', '(annexe n° 11)'
    return sum(1 for p in doc.paragraphs if _remplacer_dans_paragraphe(p, motif, rempl))


def _remplir_jetons(doc, valeurs, sequentiels=None):
    """Remplace {{X}} par la valeur (ou BLANK) et retire le placeholder « [date] ».

    `sequentiels` : {jeton: [val1, val2, …]} — pour les jetons qui se répètent et
    doivent prendre une valeur DIFFÉRENTE à chaque occurrence (ex. la commission
    paritaire, utilisée pour la ligne ouvrier PUIS la ligne employé). La 1ʳᵉ
    occurrence prend val1, la 2ᵉ val2, etc. ; au-delà on garde la dernière.
    """
    sequentiels = sequentiels or {}
    compteurs = {}

    def _rep(m):
        k = m.group(1)
        if k in sequentiels:
            vals = sequentiels[k]
            i = compteurs.get(k, 0)
            compteurs[k] = i + 1
            return str(vals[i] if i < len(vals) else vals[-1]) if vals else BLANK
        return str(valeurs.get(k, BLANK))

    for t in doc.element.body.iter(qn('w:t')):
        s = o = t.text or ''
        if '[date]' in s:
            s = s.replace('[date]', '')
        if '{{' in s:
            s = re.sub(r'\{\{(\w+)\}\}', _rep, s)
        if s != o:
            t.text = s


def _min(hhmm):
    try:
        h, m = str(hhmm).split(':')[:2]
        return int(h) * 60 + int(m)
    except Exception:
        return None


def _hhmm(minutes):
    minutes %= 1440
    return f"{minutes // 60:02d}:{minutes % 60:02d}"


def _duree(m):
    h, mm = divmod(m, 60)
    return f"{h}h{mm:02d}" if mm else f"{h}h"


def _hebdo_min(v, defaut=2280):
    """Temps plein hebdomadaire -> minutes. Accepte les notations FR et NL :
    « 36h30 », « 38h », « 37u », « 36u45 » (uur, néerlandais), « 37:30 »,
    « 36,5 », « 38 », « 37 heures »… Tolère les suffixes (heures/semaine/week)."""
    s = str(v or '').strip().lower().replace(',', '.')
    if not s:
        return defaut
    # heures + minutes : séparateur « h » (FR), « u » (NL uur) ou « : »
    m = re.match(r'^(\d+)\s*[hu:]\s*(\d{1,2})?', s)   # « 36h30 » « 38h » « 37u » « 36u45 » « 37:30 »
    if m:
        return int(m.group(1)) * 60 + int(m.group(2) or 0)
    m2 = re.match(r'^(\d+(?:\.\d+)?)', s)             # « 36.5 » « 38 » « 37 heures »
    if m2:
        return int(round(float(m2.group(1)) * 60))
    return defaut


def _cp_norm(x):
    """Clé de CP normalisée : chiffres + éventuel sous-numéro, sans « .00 » final.
    « 112 »/« 112.00 » -> « 112 » ; « 102.01 » -> « 102.01 »."""
    s = re.sub(r'[^0-9.]', '', str(x or ''))
    s = re.sub(r'\.00$', '', s)
    return s


def _cp_lookup(cp_repertoire):
    """{clé CP normalisée -> {'hebdo', 'denom'}} depuis le répertoire des CP.
    Indexe aussi la variante chiffres-seuls (ex. « 10201 ») pour tolérer les saisies."""
    out = {}
    for row in (cp_repertoire or []):
        k = _cp_norm(row.get('cp'))
        if not k:
            continue
        rec = {'hebdo': str(row.get('heures_semaine') or '').strip(),
               'denom': str(row.get('denomination') or '').strip()}
        out.setdefault(k, rec)
        dk = re.sub(r'\D', '', k)
        if dk:
            out.setdefault(dk, rec)
    return out


def _cp_info(lut, cp):
    """Retrouve {'hebdo','denom'} d'une CP dans le lookup (tolère « 112 »/« 112.00 »/« 11200 »)."""
    if not lut:
        return {}
    n = _cp_norm(cp)
    return lut.get(n) or lut.get(re.sub(r'\D', '', n)) or {}


def _jours_ouverts(jd, jf):
    """Plage de jours contiguë de jd à jf (index 0=lundi), gère le passage de semaine."""
    try:
        a, b = int(jd), int(jf)
    except (ValueError, TypeError):
        return list(range(7))
    out, i = [], a
    for _ in range(7):
        out.append(i % 7)
        if i % 7 == b % 7:
            break
        i += 1
    return out


SEUIL_PAUSE = 240   # à partir de 4h de travail on insère une pause (matin/après-midi)


def _decoupe_jour(st, work, pause):
    """Découpe une journée en matin + pause + après-midi (pause dès 4h de travail).
    Retourne (matin_de, matin_a, aprem_de, aprem_a) en minutes ; aprem_de==matin_a
    s'il n'y a pas de pause (journée très courte)."""
    if work < SEUIL_PAUSE:                              # < 4h : bloc continu
        return st, st + work, st + work, st + work
    matin = int(round((work / 2) / 30.0)) * 30          # moitié arrondie à 30 min
    matin = max(90, min(matin, work - 90))              # chaque bloc ≥ 1h30
    aprem = work - matin
    m_a = st + matin
    a_de = m_a + pause
    return st, m_a, a_de, a_de + aprem


def generer_horaires(payload, langue='FR', cible=2280, max_daily=480, min_block=180,
                     pause=30, pas=30, max_tables=600,
                     debut=None, fin=None, jour_debut=None, jour_fin=None):
    """Génère les horaires-types valides depuis les heures d'ouverture.
    Pour 5 et 6 jours de travail, avec les différentes combinaisons de jours de
    repos ET toutes les heures de début par pas de 30 min. Chaque journée intègre
    la pause si le travail dépasse 6h. Même bloc chaque jour -> 11h de repos OK.

    debut/fin/jour_debut/jour_fin : ouverture explicite (par régime) ; à défaut on
    prend celle du payload (ouverture globale)."""
    ouv_debut = debut if debut is not None else payload.get('ouverture_debut')
    ouv_fin = fin if fin is not None else payload.get('ouverture_fin')
    jd = jour_debut if jour_debut is not None else payload.get('ouverture_jour_debut', 0)
    jf = jour_fin if jour_fin is not None else payload.get('ouverture_jour_fin', 6)
    d, f = _min(ouv_debut), _min(ouv_fin)
    if d is None or f is None:
        return []
    ouverts = _jours_ouverts(jd, jf)
    if len(ouverts) < 5:
        return []
    win_s = d
    win_e = f if f > d else f + 1440           # ouverture de nuit (ex. 05:00->03:00)
    noms = JOURS_NL if langue == 'NL' else JOURS_FR
    schedules = []

    from itertools import combinations

    def arrangements():
        # 6 jours puis 5 jours ; repos = TOUTES les combinaisons possibles de jours
        # de repos parmi les jours d'ouverture (pas seulement consécutives).
        n_open = len(ouverts)
        for n_work, label in ((6, '6 jours'), (5, '5 jours')):
            if n_open < n_work:
                continue
            n_off = n_open - n_work
            combos = [()] if n_off <= 0 else combinations(ouverts, n_off)
            for off_open in combos:
                work = [x for x in ouverts if x not in off_open]
                if len(work) != n_work:
                    continue
                off_disp = [d for d in range(7) if d not in work]  # tous les jours non prestés
                lab = label if langue != 'NL' else label.replace('jours', 'dagen')
                yield work, off_disp, lab

    win_len = win_e - win_s

    for work, off, mode in arrangements():
        n = len(work)
        plein = cible // n                                    # min/jour requis pour le temps plein
        span_plein = plein + (pause if plein >= SEUIL_PAUSE else 0)
        # « complet » = l'horaire atteint réellement le temps plein (rentre dans la
        # fenêtre d'ouverture ET respecte le max journalier légal).
        complet = (plein <= max_daily) and (span_plein <= win_len)
        if complet:
            daily = plein
        elif (win_len - pause) >= min_block:
            daily = min(win_len - pause, max_daily)            # remplit la fenêtre (SOUS le temps plein)
        else:
            continue
        span = daily + (pause if daily >= SEUIL_PAUSE else 0)
        workset = set(work)
        total = daily * n
        approx = '' if total == cible else (
            f" · {_duree(total)}/sem" if langue != 'NL' else f" · {_duree(total)}/week")

        def ajoute(st):
            m_de, m_a, a_de, a_a = _decoupe_jour(st, daily, pause)
            # TOUS les 7 jours : jour presté = horaire, sinon repos (None)
            lignes = [((day, m_de, m_a, a_de, a_a) if day in workset
                       else (day, None, None, None, None)) for day in range(7)]
            repos = ', '.join(noms[o] for o in off)
            titre = (f"{mode} — {'repos' if langue != 'NL' else 'rust'} {repos}"
                     f" · {'début' if langue != 'NL' else 'start'} {_hhmm(st)}{approx}")
            schedules.append({'titre': titre, 'lignes': lignes, 'off': off,
                              'total': total, 'complet': complet,
                              'pause': pause if daily >= SEUIL_PAUSE else 0})

        st, last = win_s, None
        while st + span <= win_e:
            ajoute(st); last = st
            if len(schedules) >= max_tables:
                return schedules
            st += pas
        # dernier horaire calé sur la fermeture (pour couvrir jusqu'à l'heure de fin)
        close_st = win_e - span
        if last is not None and close_st > last:
            ajoute(close_st)
            if len(schedules) >= max_tables:
                return schedules
    return schedules


def _bordures(tbl):
    """Applique des bordures fines à un tableau (sans dépendre d'un style nommé)."""
    from docx.oxml import OxmlElement
    tblPr = tbl._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), '4')
        e.set(qn('w:space'), '0'); e.set(qn('w:color'), 'BFC4D6')
        borders.append(e)
    tblPr.append(borders)


def _ajouter_annexe_horaires(doc, schedules, langue='FR', titre_section=None):
    from docx.shared import Pt, RGBColor
    noms = JOURS_NL if langue == 'NL' else JOURS_FR
    ACCENT = RGBColor(0x1c, 0x22, 0x44)
    doc.add_page_break()
    h = doc.add_paragraph()
    entete = titre_section or ('ANNEXE — HORAIRES DE TRAVAIL POSSIBLES' if langue != 'NL'
                               else 'BIJLAGE — MOGELIJKE UURROOSTERS')
    r = h.add_run(entete)
    r.bold = True; r.font.size = Pt(14); r.font.color.rgb = ACCENT
    intro = doc.add_paragraph()
    ri = intro.add_run(
        "Horaires-types compatibles avec les heures d'ouverture. "
        "L'employeur communique l'horaire applicable au travailleur (à valider juridiquement)."
        if langue != 'NL' else
        "Mogelijke uurroosters binnen de openingsuren. Juridisch te valideren.")
    ri.italic = True; ri.font.size = Pt(9)
    thd = (['Jour', 'Matin', 'Pause', 'Après-midi', 'Durée'] if langue != 'NL'
           else ['Dag', 'Voormiddag', 'Pauze', 'Namiddag', 'Duur'])
    for sc in schedules:
        p = doc.add_paragraph()
        rp = p.add_run(sc['titre']); rp.bold = True; rp.font.size = Pt(10.5); rp.font.color.rgb = ACCENT
        p.paragraph_format.space_before = Pt(9); p.paragraph_format.space_after = Pt(2)
        tbl = doc.add_table(rows=1, cols=5)
        _bordures(tbl)
        for i, t in enumerate(thd):
            tbl.rows[0].cells[i].paragraphs[0].add_run(t).bold = True
        repos_txt = 'Repos' if langue != 'NL' else 'Rust'
        for day, m_de, m_a, a_de, a_a in sc['lignes']:
            c = tbl.add_row().cells
            c[0].paragraphs[0].add_run(noms[day].capitalize())
            if m_de is None:                      # jour de repos : ligne vide (visible)
                c[1].paragraphs[0].add_run(repos_txt)
                for k in (2, 3, 4):
                    c[k].paragraphs[0].add_run('—')
                continue
            c[1].paragraphs[0].add_run(f"{_hhmm(m_de)} – {_hhmm(m_a)}")
            c[2].paragraphs[0].add_run(f"{_hhmm(m_a)} – {_hhmm(a_de)}" if a_de > m_a else '—')
            c[3].paragraphs[0].add_run(f"{_hhmm(a_de)} – {_hhmm(a_a)}" if a_a > a_de else '—')
            c[4].paragraphs[0].add_run(_duree((m_a - m_de) + (a_a - a_de)))
        tot = doc.add_paragraph()
        pz = (f" · pause {sc['pause']} min/jour" if sc.get('pause') else '')
        rt = tot.add_run(('Total : ' if langue != 'NL' else 'Totaal: ') + _duree(sc['total']) + ' / '
                         + ('semaine' if langue != 'NL' else 'week') + pz)
        rt.font.size = Pt(9); rt.italic = True


def _remplir_cameras(doc, nb, lieux, lang='FR'):
    """Annexe 7 (Caméras) : « comporte … caméra(s) … aux endroits suivants : … ».
    Ces blancs sont des runs de « ……… » (pas des jetons). On remplit le 1er blanc
    (nombre) et le 2e (emplacements) du paragraphe caméras, sans toucher au reste
    (préserve la mise en page). Défaut : 0 caméra / « Néant »."""
    kw = 'caméra' if lang != 'NL' else 'camera'
    mk = 'comporte' if lang != 'NL' else 'bestaat uit'

    def isblank(s):
        t = (s or '').strip()
        return bool(t) and set(t) <= set('….') and len(t) >= 2

    for p in doc.paragraphs:
        runs = p.runs
        full = ''.join(r.text or '' for r in runs).lower()
        if kw not in full or mk not in full:
            continue
        blanks = [i for i, r in enumerate(runs) if isblank(r.text)]
        if not blanks:
            return False
        # remplacement par FONCTION : le texte utilisateur (nb/lieux) n'est jamais
        # réinterprété comme motif de remplacement (\1, \g<>, &, \ … restent littéraux).
        # Nettoyage résiduel : on retire les « … » ou 2+ points, mais PAS un point isolé
        # (qui pourrait terminer une phrase légitime).
        resid = r'^\s*(?:…+|\.{2,})'
        i0 = blanks[0]
        runs[i0].text = re.sub(r'[….]+', lambda _m: str(nb), runs[i0].text, count=1)
        if i0 + 1 < len(runs):                       # ex. NL « ..camera » -> « camera »
            runs[i0 + 1].text = re.sub(resid, '', runs[i0 + 1].text or '')
        if len(blanks) >= 2:
            i1 = blanks[1]
            runs[i1].text = re.sub(r'[….]+', lambda _m: str(lieux), runs[i1].text, count=1)
            if i1 + 1 < len(runs):
                runs[i1 + 1].text = re.sub(resid, '', runs[i1 + 1].text or '')
        return True
    return False


def _remplir_placeholders_litteraux(doc, valeurs):
    """Les modèles FR ET NL contiennent, à quelques endroits, du TEXTE LITTÉRAL
    « Nom (Em) », « No ONSS (Em) » et — annexe 5 du NL — « Rue (Em), No de la maison
    (Em) Code postal (Em) Localité (Em) » qui n'ont jamais été convertis en jetons au
    prétraitement. Sans ça, ces noms de champs s'affichent bruts. On les remplit ici.
    Chaque branche est un no-op là où le littéral n'existe pas."""
    nom = valeurs.get('Nom_Em', BLANK)
    onss = valeurs.get('No_ONSS_Em', BLANK)
    expl = valeurs.get('sieges_exploitation', BLANK)
    simples = {'Nom (Em)': nom, 'No ONSS (Em)': onss}
    for p in doc.paragraphs:
        full = ''.join(r.text or '' for r in p.runs)
        # Annexe 5 : le paragraphe ne contient QUE le placeholder d'adresse d'exploitation
        if 'Rue (Em)' in full and 'No de la maison (Em)' in full:
            done = False
            for r in p.runs:
                if not done and 'Rue (Em)' in (r.text or ''):
                    r.text, done = str(expl), True
                else:
                    r.text = ''
            continue
        # Remplacements simples (le littéral vit dans un seul run, on préserve le reste)
        for r in p.runs:
            t = r.text or ''
            if not t:
                continue
            for lit, val in simples.items():
                if lit in t:
                    t = t.replace(lit, str(val))
            if t != r.text:
                r.text = t

    # Cadre NL : « ……….u en ………u » résiduel collé après la plage horaire -> on nettoie
    for p in doc.paragraphs:
        full = ''.join(r.text or '' for r in p.runs)
        if 'arbeidsprestaties kunnen worden vastgesteld' in full and '…' in full:
            apres = False
            for r in p.runs:
                t = r.text or ''
                if not apres:
                    if 'tussen' in t or 'vastgesteld' in t:
                        apres = True
                    continue
                if t and set(t) <= set('….uen '):   # run purement résiduel (… . u e n espace)
                    r.text = ''
            break

    # Annexe 4 bien-être du modèle NL : pas de jetons, juste des labels « … is (zijn) : »
    # (le modèle FR utilise des jetons -> ces ancres NL n'y existent pas, no-op).
    _ajouter_apres_label(doc, 'preventieadviseur(s) is (zijn)', valeurs.get('personne_confiance'))
    _ajouter_apres_label(doc, 'Geweld, pesterijen en ongewenst seksueel gedrag op het werk is (zijn)',
                         valeurs.get('harcelement'))


def _ajouter_max_hebdo(doc, lang):
    """Ajoute la ligne « Maximum 50 heures par semaine … » juste après la durée
    journalière maximale dans le cadre (Article 10 §2). La loi du 1/6/2026 exige cette
    mention du max hebdomadaire ; le modèle n'a pas de jeton pour elle -> on l'insère."""
    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph
    if lang != 'NL':
        ancre, prefixe = 'durée journalière maximale de travail', '-\t'
        texte = 'Maximum 50 heures par semaine selon les limites journalières maximales.'
    else:
        ancre, prefixe = 'maximale dagelijkse arbeidsduur', ''
        texte = 'Maximum 50 uren per week volgens de maximale dagelijkse grenzen.'
    for p in doc.paragraphs:
        full = ''.join(r.text or '' for r in p.runs)
        if ancre in full:
            if '50' in full:                       # déjà présent -> ne pas dupliquer
                return
            new_p = OxmlElement('w:p')
            p._p.addnext(new_p)
            np = Paragraph(new_p, p._parent)
            try:
                np.style = p.style                 # même mise en forme que la ligne au-dessus
            except Exception:
                pass
            np.add_run(prefixe + texte)
            return


def _ajouter_apres_label(doc, ancre, valeur):
    """Ajoute `valeur` à la fin du 1er paragraphe contenant `ancre` (label NL qui finit
    par « : »). No-op si `valeur` est vide/BLANK, si l'ancre est absente, ou si la valeur
    est déjà présente."""
    if not valeur or valeur == BLANK:
        return
    valeur = str(valeur)
    for p in doc.paragraphs:
        full = ''.join(r.text or '' for r in p.runs)
        if ancre in full:
            if valeur in full:
                return
            p.add_run(' ' + valeur)
            return


def build_reglement(payload, identity=None, template_bytes=None, model_bytes=None,
                    repertoire=None, cp_repertoire=None):
    """Remplit le modèle officiel et renvoie les bytes du .docx.
    template_bytes : le modèle FR/NL (jetons {{...}}) depuis Supabase Storage.
    model_bytes    : (optionnel) horaire sectoriel à ajouter en fin de document.
    repertoire     : liste d'institutions (adresses) pour remplir Art. 2 & 66.
    cp_repertoire  : liste des commissions paritaires (n° + dénomination + temps plein)
                     pour compléter la dénomination de la CP en Article 2.
    """
    if not template_bytes:
        raise ValueError("Modèle de règlement introuvable (pas encore hébergé).")
    lang = 'NL' if str(payload.get('reglement_langue') or 'FR').upper() == 'NL' else 'FR'
    lut = _cp_lookup(cp_repertoire)
    doc = Document(io.BytesIO(template_bytes))
    # Commission paritaire : la MÊME balise sert pour les 2 lignes du modèle
    # (ouvrier + employé) -> remplacement positionnel (1ʳᵉ occ. = 1er régime/CP,
    # 2ᵉ occ. = 2e régime/CP). Si des régimes explicites sont fournis, on les utilise.
    regs = [r for r in (payload.get('regimes') or []) if isinstance(r, dict)] \
        if isinstance(payload.get('regimes'), list) else []
    if regs:
        def _cpn(i):
            # _cp_norm : « 140.03 » doit s'imprimer « 140.03 », pas « 14003 ».
            return _cp_norm(regs[i].get('cp')) if i < len(regs) else ''
        def _den(i):
            # Dénomination OFFICIELLE (liste des Fonds de sécurité d'existence du SPF
            # Emploi). On n'utilise PAS regs[i]['label'] : c'est le libellé libre de
            # l'écran (« Car-wash », « Nettoyage »), un surnom interne — l'imprimer ici
            # donnait « Dénomination : Car-wash » au lieu de la dénomination légale.
            if i >= len(regs):
                return ''
            return (_fse_info(regs[i].get('cp'), lang).get('denomination')
                    or _cp_info(lut, regs[i].get('cp')).get('denom', '')
                    or (regs[i].get('denomination') or '').strip())
        cp_ouv = _cpn(0) or BLANK
        cp_emp = _cpn(1) or BLANK
        den_ouv = _den(0) or BLANK
        den_emp = _den(1) or BLANK
    else:
        def _den_simple(cp, saisi):
            return ((_fse_info(cp, lang).get('denomination')
                     or _cp_info(lut, cp).get('denom', '')
                     or (saisi or '').strip()) or BLANK)
        cp_ouv = _cp_norm(payload.get('commission_paritaire')) or BLANK
        cp_emp = _cp_norm(payload.get('cp_employe')) or BLANK
        den_ouv = _den_simple(payload.get('commission_paritaire'),
                              payload.get('cp_ouvrier_denomination'))
        den_emp = _den_simple(payload.get('cp_employe'),
                              payload.get('cp_employe_denomination'))
    valeurs = _valeurs(payload, identity, repertoire)
    sequentiels = {
        'Commission_paritaire_Em': [cp_ouv, cp_emp],
        'Commission_paritaire_Em_v1': [den_ouv, den_emp],
    }
    if lang != 'NL':
        # Le modèle FR réutilise {{Nom_1_institution_Inst}} pour le point 6 (SEPPT)
        # PUIS le point 8 (Bureau des contributions directes) : sans remplacement
        # positionnel, le nom du SEPPT s'imprime aussi comme bureau des contributions.
        # 1ʳᵉ occurrence = SEPPT, 2ᵉ = blanc à compléter. (Le NL n'a pas ce doublon.)
        sequentiels['Nom_1_institution_Inst'] = [
            valeurs.get('Nom_1_institution_Inst', BLANK), BLANK]
    _remplir_jetons(doc, valeurs, sequentiels)
    # Annexe 7 — caméras de surveillance (défaut 0 si le client n'en a pas)
    nb_cam = str(payload.get('nombre_cameras') or payload.get('cameras') or '').strip() or '0'
    lieux_cam = (payload.get('cameras_emplacement') or '').strip() or (
        ('Néant' if lang != 'NL' else 'Geen') if nb_cam in ('0', '') else BLANK)
    _remplir_cameras(doc, nb_cam, lieux_cam, lang)
    # FR + NL : remplit les placeholders littéraux résiduels (Nom, No ONSS, annexe 5 NL)
    _remplir_placeholders_litteraux(doc, valeurs)
    # Cadre horaire (loi 1/6/2026) : ajoute le max hebdomadaire (50h) après le max journalier
    _ajouter_max_hebdo(doc, lang)
    # Article 4 : le renvoi doit viser l'annexe 11 (accusé de réception), pas la 10
    _corriger_renvoi_annexe(doc, lang)
    # NB : les horaires générés vont dans un DOCUMENT SÉPARÉ (generer_doc_horaires),
    # plus le règlement lui-même, car ils peuvent faire des centaines de pages.

    if model_bytes:
        # Ajoute aussi l'horaire sectoriel choisi. Certains « modèles » sont en
        # réalité des fichiers Excel (non joignables) -> on n'échoue PAS.
        try:
            from docxcompose.composer import Composer
            doc.add_page_break()
            buf = io.BytesIO(); doc.save(buf); buf.seek(0)
            comp = Composer(Document(buf))
            comp.append(Document(io.BytesIO(model_bytes)))
            out = io.BytesIO(); comp.save(out)
            return out.getvalue()
        except Exception as e:
            print(f"[REGLEMENT] annexe modèle ignorée (non joignable) : {e}")

    out = io.BytesIO(); doc.save(out)
    return out.getvalue()


def _regimes_du_payload(payload, cp_repertoire=None):
    """Normalise les régimes de travail à générer. Chaque régime = une CP avec SON
    temps plein et SES propres heures d'ouverture -> une section d'horaires dédiée.

    - Si payload['regimes'] est fourni (liste), on l'utilise ; le temps plein et le
      libellé manquants sont complétés depuis le répertoire des CP (sinon 38h/sem).
    - Sinon rétro-compat : ouvriers (CP principale + ouverture globale) et,
      si une 2e CP/temps plein employés existe, une section employés (même ouverture).
    """
    lut = _cp_lookup(cp_repertoire)

    def _mxj(v):
        return (_min(f"{v}:00") if str(v or '').isdigit() else 480) or 480

    def _cible(cp, hebdo):
        h = hebdo if str(hebdo or '').strip() else _cp_info(lut, cp).get('hebdo')
        return _hebdo_min(h, 2280)

    def _lib(cp, label):
        return (label or '').strip() or _cp_info(lut, cp).get('denom', '')

    regs = payload.get('regimes')
    items = []
    if isinstance(regs, list) and regs:
        for rg in regs:
            if not isinstance(rg, dict):            # entrée malformée -> ignorée
                continue
            deb, fin = rg.get('ouverture_debut'), rg.get('ouverture_fin')
            if not (deb and fin):
                continue
            cp = str(rg.get('cp') or '').strip()
            hebdo = rg.get('hebdo') if rg.get('hebdo') not in (None, '') else rg.get('heures_semaine')
            items.append({
                'label': _lib(cp, rg.get('label')),
                'cp': cp,
                'cible': _cible(cp, hebdo),
                'debut': deb, 'fin': fin,
                'jd': rg.get('jour_debut', rg.get('ouverture_jour_debut', 0)),
                'jf': rg.get('jour_fin', rg.get('ouverture_jour_fin', 6)),
                'mx': _mxj(rg.get('max_journalier')),
            })
        return items

    # Rétro-compat (ancien modèle ouvrier/employé, ouverture globale unique)
    if not (payload.get('ouverture_debut') and payload.get('ouverture_fin')):
        return []
    base = {'debut': payload.get('ouverture_debut'), 'fin': payload.get('ouverture_fin'),
            'jd': payload.get('ouverture_jour_debut', 0), 'jf': payload.get('ouverture_jour_fin', 6),
            'mx': _mxj(payload.get('max_journalier'))}
    cp_o = payload.get('commission_paritaire')
    items.append({**base, 'label': _lib(cp_o, 'ouvriers'), 'cp': cp_o,
                  'cible': _cible(cp_o, payload.get('heures_ouvrier'))})
    cp_e = str(payload.get('cp_employe') or '').strip()
    h_e = payload.get('heures_employe')
    if cp_e or h_e:
        items.append({**base, 'label': _lib(cp_e or cp_o, 'employés'), 'cp': cp_e or cp_o,
                      'cible': _cible(cp_e or cp_o, h_e)})
    return items


def _ajouter_avertissement_ouverture(doc, it, cible, max_total, langue, titre):
    """Écrit, à la place des tableaux, un avertissement clair quand les heures
    d'ouverture saisies ne permettent pas d'atteindre le temps plein (ex. 38h dans
    une fenêtre de 7h sur 5 jours). Évite de sortir des horaires « 38h » à 32h30."""
    from docx.shared import Pt, RGBColor
    AMBRE = RGBColor(0xb4, 0x53, 0x09)
    noms = JOURS_NL if langue == 'NL' else JOURS_FR
    ouv = _jours_ouverts(it['jd'], it['jf'])
    jours_txt = f"{noms[ouv[0]]}–{noms[ouv[-1]]}"
    deb, fin = it['debut'], it['fin']
    doc.add_page_break()
    h = doc.add_paragraph()
    rh = h.add_run(titre); rh.bold = True; rh.font.size = Pt(14); rh.font.color.rgb = AMBRE
    p = doc.add_paragraph()
    if langue != 'NL':
        msg = (f"⚠ Les heures d'ouverture saisies ({deb}–{fin}, {jours_txt} = {len(ouv)} jours "
               f"d'ouverture) ne permettent pas d'atteindre le temps plein de {_duree(cible)}/semaine. "
               f"Maximum atteignable dans cette fenêtre : environ {_duree(max_total)}/semaine. "
               f"→ Élargissez la plage horaire ou ajoutez des jours d'ouverture, puis régénérez.")
    else:
        msg = (f"⚠ De ingevoerde openingsuren ({deb}–{fin}, {jours_txt} = {len(ouv)} openingsdagen) "
               f"laten niet toe om het voltijds uurrooster van {_duree(cible)}/week te bereiken. "
               f"Maximaal haalbaar in dit venster: ongeveer {_duree(max_total)}/week. "
               f"→ Verruim de openingsuren of voeg openingsdagen toe en genereer opnieuw.")
    rp = p.add_run(msg); rp.font.size = Pt(11); rp.font.color.rgb = AMBRE


def generer_doc_horaires(payload, identity=None, cp_repertoire=None):
    """Document Word SÉPARÉ contenant tous les horaires-types (5j + 6j, toutes les
    combinaisons de jours de repos, tous les débuts par 30 min), UNE SECTION PAR
    RÉGIME (chaque CP avec son temps plein et sa propre ouverture).
    Renvoie les bytes, ou None si aucun horaire."""
    lang = 'NL' if str(payload.get('reglement_langue') or 'FR').upper() == 'NL' else 'FR'
    items = _regimes_du_payload(payload, cp_repertoire)
    if not items:
        return None

    from docx.shared import Pt, RGBColor
    doc = Document()
    stl = doc.styles['Normal']; stl.font.name = 'Calibri'; stl.font.size = Pt(10)
    ti = doc.add_paragraph()
    rt = ti.add_run(('HORAIRES DE TRAVAIL — ' if lang != 'NL' else 'UURROOSTERS — ')
                    + (identity or {}).get('nom_societe', ''))
    rt.bold = True; rt.font.size = Pt(16); rt.font.color.rgb = RGBColor(0x1c, 0x22, 0x44)

    total_tables = 0
    for it in items:
        try:
            scheds = generer_horaires(payload, langue=lang, cible=it['cible'], max_daily=it['mx'],
                                      max_tables=1500, debut=it['debut'], fin=it['fin'],
                                      jour_debut=it['jd'], jour_fin=it['jf'])
        except Exception as e:
            print(f"[REGLEMENT] horaires {it.get('label') or it.get('cp')} échoués : {e}"); continue
        if not scheds:
            continue
        cpnum = re.sub(r'[^0-9.]', '', str(it['cp'])) if it['cp'] else ''
        cph = f" (CP {cpnum})" if cpnum else ''
        base = (it['label'] or (f"CP {cpnum}" if cpnum else 'Horaires')).upper()
        titre = (f"{base}{cph} — {_duree(it['cible'])}/semaine" if lang != 'NL'
                 else f"{base}{cph} — {_duree(it['cible'])}/week")
        # On ne garde QUE les horaires qui atteignent réellement le temps plein.
        complets = [s for s in scheds if s.get('complet')]
        if complets:
            _ajouter_annexe_horaires(doc, complets, lang, titre_section=titre)
            total_tables += len(complets)
        else:
            # Aucun horaire n'atteint le temps plein -> l'ouverture est trop étroite :
            # on l'écrit clairement au lieu de tableaux trompeurs sous le mauvais total.
            max_total = max((s['total'] for s in scheds), default=0)
            _ajouter_avertissement_ouverture(doc, it, it['cible'], max_total, lang, titre)
            total_tables += 1
    if total_tables == 0:
        return None
    out = io.BytesIO(); doc.save(out)
    return out.getvalue()
